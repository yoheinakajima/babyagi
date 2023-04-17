import random
import string
from docx import Document
import os


class LoggerAgent:
    def __init__(self):
        self.document = Document()
        subfolder = "output"
        if not os.path.exists(subfolder):
            os.makedirs(subfolder)
        self.file_name = f"{subfolder}/result_{self.generate_random_string()}.doc"

    def log(self, message):
        paragraph = self.document.add_paragraph(message)
        paragraph.add_run("\n")  # add a new line after each message
        self.document.save(self.file_name)

    def close(self):
        self.document.save(self.file_name)
        #self.document.close()

    def generate_random_string(self):
        """Generate a random string of length 6"""
        return ''.join(random.choices(string.ascii_lowercase + string.digits, k=6))
